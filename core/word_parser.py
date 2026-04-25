from pathlib import Path
from typing import List, Dict
from docx import Document


class WordParser:
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.document = None
        self.raw_text = ""
        self.paragraphs = []
        
    def load(self) -> bool:
        try:
            self.document = Document(str(self.file_path))
            self.raw_text = "\n".join([para.text for para in self.document.paragraphs])
            self.paragraphs = [para.text for para in self.document.paragraphs if para.text.strip()]
            return True
        except Exception as e:
            print(f"Error loading Word file: {e}")
            return False
    
    def extract_content(self) -> Dict[str, any]:
        content = {
            "title": self._extract_title(),
            "sections": [],
            "lists": [],
            "questions": [],
            "full_text": self.raw_text
        }
        
        current_section = None
        current_list = []
        in_list = False
        
        for para in self.paragraphs:
            para_text = para.strip()
            
            if self._is_question(para_text):
                content["questions"].append({
                    "question": para_text,
                    "answer": ""
                })
            elif self._is_list_item(para_text):
                if not in_list:
                    in_list = True
                    current_list = []
                current_list.append(self._clean_list_item(para_text))
            else:
                if in_list and current_list:
                    content["lists"].append(current_list)
                    in_list = False
                    current_list = []
                
                if para_text and len(para_text) > 10:
                    content["sections"].append({
                        "text": para_text,
                        "type": self._detect_section_type(para_text)
                    })
        
        if in_list and current_list:
            content["lists"].append(current_list)
        
        return content
    
    def _extract_title(self) -> str:
        if self.paragraphs:
            return self.paragraphs[0].strip()
        return ""
    
    def _is_question(self, text: str) -> bool:
        return "?" in text or text.lower().startswith(("what", "how", "why", "when", "where", "who", "nasıl", "ne", "neden"))
    
    def _is_list_item(self, text: str) -> bool:
        return text.strip().startswith(("-", "•", "*", "1.", "2.", "3.", "a)", "b)", "c)"))
    
    def _clean_list_item(self, text: str) -> str:
        cleaned = text.strip()
        for prefix in ["-", "•", "*", "1.", "2.", "3.", "a)", "b)", "c)"]:
            if cleaned.startswith(prefix):
                cleaned = cleaned[len(prefix):].strip()
        return cleaned
    
    def _detect_section_type(self, text: str) -> str:
        text_lower = text.lower()
        if any(word in text_lower for word in ["örnek", "example", "mesela"]):
            return "example"
        elif any(word in text_lower for word in ["ipucu", "tip", "not"]):
            return "tip"
        elif any(word in text_lower for word in ["kural", "rule", "grammar"]):
            return "rule"
        else:
            return "general"
